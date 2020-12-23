#! python3

### DOCX
#pip install python-docx
import docx

d = docx.Document('c:\\Users\\Aristos\\Desktop\\Aristos\\test.docx') #return a document object
d.paragraphs #returns a list of all the paragraphs
d.paragraphs[0] 
d.paragraphs[0].text

p = d.paragraphs[1] #select the second paragraph
p.runs #a run is defined as the scan of every character until you encounter a change in style e.g. bold or any kind of "different" text.
#so p.runs returns a list with all the runs

p.runs[3].text
p.runs[3].bold #returns bool

p.runs[0].bold = True# i make the first word BOLD. The change is not in effect until the file is saved.
d.save('c:\\Users\\Aristos\\Desktop\\Aristos\\test.docx') #now i save it (with the same name), so its like hiting CTRL+S

p.style

new_doc = docx.Document() #creates a blank document locally, not in the hard drive.
d.add_paragraph('Hello this is the 1st paragraph')
new_doc.save('c:\\Users\\Aristos\\Desktop\\Aristos\\new_test.docx')
p = d.paragraph[0]
p.add_run('This is a new run.')
p.runs
p.runs[0].bold = True

#return 
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
